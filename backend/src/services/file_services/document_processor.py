# app/services/file_service.py
import os
import logging
import io
from fastapi import UploadFile, HTTPException
from typing import Tuple
from config.logger import setup_logging

# Text extraction imports
import PyPDF2
import pdfplumber
from docx import Document


# Setup logging
setup_logging() 
logger = logging.getLogger(__name__)


class DocumentProcessor:
    def __init__(self):
        pass
            
    
    async def get_file_type(self, filename: str) -> str:
        """Extract file type from filename."""
        logger.debug(f"Extracting file type from filename: {filename}")
        extension = os.path.splitext(filename)[1].lower()
        
        if extension == ".pdf":
            logger.debug(f"File type detected: pdf")
            return "pdf"
        elif extension == ".docx":
            logger.debug(f"File type detected: docx")
            return "docx"
        else:
            logger.warning(f"Unsupported file type: {extension}")
            raise HTTPException(
                status_code=400, 
                detail="Unsupported file type. Supported types: pdf, docx"
            )
    
    async def extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF file content.
        
        Args:
            file_content (bytes): PDF file content
            
        Returns:
            str: Extracted text content
        """
        logger.info("Extracting text from PDF file")
        try:
            # Try with pdfplumber first (better for complex layouts)
            with pdfplumber.open(io.BytesIO(file_content)) as pdf:
                text_content = []
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text_content.append(page_text)
                
                if text_content:
                    extracted_text = "\n\n".join(text_content)
                    logger.info(f"Successfully extracted {len(extracted_text)} characters using pdfplumber")
                    return extracted_text
            
            # Fallback to PyPDF2 if pdfplumber fails
            logger.info("Falling back to PyPDF2 for PDF extraction")
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_content = []
            
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            
            extracted_text = "\n\n".join(text_content)
            logger.info(f"Successfully extracted {len(extracted_text)} characters using PyPDF2")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from PDF: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text from PDF: {str(e)}"
            )
    
    async def extract_text_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file content.
        
        Args:
            file_content (bytes): DOCX file content
            
        Returns:
            str: Extracted text content
        """
        logger.info("Extracting text from DOCX file")
        try:
            doc = Document(io.BytesIO(file_content))
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(" | ".join(row_text))
            
            extracted_text = "\n\n".join(text_content)
            logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
            return extracted_text
            
        except Exception as e:
            logger.error(f"Failed to extract text from DOCX: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=500,
                detail=f"Failed to extract text from DOCX: {str(e)}"
            )
    
    async def extract_text_content(self, file_content: bytes, file_type: str) -> str:
        """
        Extract text content from file based on file type.
        
        Args:
            file_content (bytes): File content as bytes
            file_type (str): Type of file (pdf, docx, pptx, txt)
            
        Returns:
            str: Extracted and formatted text content
        """
        logger.info(f"Extracting text content for file type: {file_type}")
        
        if file_type == "pdf":
            return await self.extract_text_from_pdf(file_content)
        elif file_type == "docx":
            return await self.extract_text_from_docx(file_content)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type for text extraction: {file_type}"
            )
    
    async def extract_text_from_upload(self, file: UploadFile) -> Tuple[str, str]:
        """
        Extract text content from uploaded file.
        
        Args:
            file (UploadFile): The uploaded file
            
        Returns:
            Tuple[str, str]: (extracted_text, file_type)
        """
        logger.info(f"Extracting text from uploaded file: {file.filename}")
        file_type = await self.get_file_type(file.filename)
        
        # Read file content
        file_content = await file.read()
        logger.debug(f"Read file content, size: {len(file_content)} bytes")
        
        # Extract text content
        extracted_text = await self.extract_text_content(file_content, file_type)
        
        # Reset file position for potential future reads
        await file.seek(0)
        
        return extracted_text, file_type